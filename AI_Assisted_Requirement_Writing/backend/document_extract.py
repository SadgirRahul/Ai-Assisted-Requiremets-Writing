"""
document_extract.py
Extract plain text from PDF and Word (.docx) documents for the requirements pipeline.
"""
import os
from typing import List

import PyPDF2
from docx import Document


def _extract_text_from_pdf(path: str) -> str:
    text_parts: List[str] = []
    with open(path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def extract_text(path: str) -> str:
    """
    Extract plain text from a supported document.

    Supported extensions: .pdf, .docx
    Legacy .doc (Word 97-2003) is not supported.
    """
    if not path or not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_text_from_pdf(path)
    if ext == ".docx":
        return _extract_text_from_docx(path)

    raise ValueError(
        f"Unsupported file type '{ext}'. Supported types: .pdf, .docx"
    )
